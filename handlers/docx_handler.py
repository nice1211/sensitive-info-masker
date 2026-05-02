"""
Word 文档处理器 (.docx)
遍历段落和表格单元格进行脱敏/还原
"""
from docx import Document
from core.detector import SensitiveDetector
from core.masker import Masker


def mask_docx_file(filepath, detector: SensitiveDetector, items=None):
    """
    对 docx 文件执行脱敏
    items: 可选，前端传来的已筛选检测项列表。为 None 时自动检测。
    返回: (Document 对象, mapping)
    """
    doc = Document(filepath)
    masker = Masker()

    # 处理段落
    for para in doc.paragraphs:
        _mask_paragraph(para, detector, masker, items)

    # 处理表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _mask_paragraph(para, detector, masker, items)

    # 处理页眉页脚
    for section in doc.sections:
        for header_footer in [section.header, section.footer]:
            if header_footer is not None:
                for para in header_footer.paragraphs:
                    _mask_paragraph(para, detector, masker, items)

    return doc, masker.get_mapping()


def unmask_docx_file(filepath, mapping):
    """
    对 docx 文件执行还原
    返回: Document 对象
    """
    doc = Document(filepath)

    for para in doc.paragraphs:
        _unmask_paragraph(para, mapping)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _unmask_paragraph(para, mapping)

    for section in doc.sections:
        for header_footer in [section.header, section.footer]:
            if header_footer is not None:
                for para in header_footer.paragraphs:
                    _unmask_paragraph(para, mapping)

    return doc


def _mask_paragraph(para, detector, masker, items=None):
    """对单个段落执行脱敏（保留格式）"""
    full_text = para.text
    if not full_text.strip():
        return

    if items is not None:
        # 使用前端筛选的项：只匹配 value 在当前段落中出现的
        import re
        detections = []
        for item in items:
            for m in re.finditer(re.escape(item['value']), full_text):
                detections.append({
                    'type': item['type'],
                    'value': item['value'],
                    'start': m.start(),
                    'end': m.end(),
                })
        detections.sort(key=lambda x: x['start'])
    else:
        detections = detector.detect(full_text)

    if not detections:
        return

    masked_text, _ = masker.mask_text(full_text, detections)

    # 保留第一个 run 的格式，清除其余 run
    if para.runs:
        first_run = para.runs[0]
        for run in para.runs:
            run.text = ''
        first_run.text = masked_text


def _unmask_paragraph(para, mapping):
    """对单个段落执行还原"""
    for run in para.runs:
        if run.text:
            run.text = Masker.unmask_text(run.text, mapping)
