"""
敏感信息加密工具 - Flask 主程序
"""
import os
import re
import time
from datetime import datetime
from flask import Flask, request, jsonify, send_file, send_from_directory

import config
from core.detector import SensitiveDetector
from core.masker import Masker
from core.crypto import CryptoManager
from handlers.text_handler import mask_text_file, unmask_text_file, save_text
from handlers.docx_handler import mask_docx_file, unmask_docx_file
from handlers.xlsx_handler import mask_xlsx_file, unmask_xlsx_file

app = Flask(__name__, static_folder='static', static_url_path='')

# 初始化组件
detector = SensitiveDetector(config.CUSTOM_NAMES_FILE)
crypto = CryptoManager(config.KEY_FILE)

TEXT_EXTS = {'.txt', '.md', '.csv'}
DOCX_EXTS = {'.docx'}
XLSX_EXTS = {'.xlsx'}


def _ts():
    return datetime.now().strftime('%Y%m%d_%H%M%S')


def _safe_filename(filename):
    """安全化文件名：保留中文、字母、数字、下划线、点、短横线，去掉路径分隔符等危险字符"""
    # 去掉目录部分，只保留文件名
    filename = filename.replace('\\', '/').split('/')[-1]
    # 只保留安全字符
    filename = re.sub(r'[^\w\u4e00-\u9fff.\-]', '_', filename)
    # 防止空文件名
    if not filename or filename.startswith('.'):
        filename = 'unnamed' + filename
    return filename


# ─── 页面 ──────────────────────────────────────────────────────

@app.route('/')
def index():
    return send_from_directory('static', 'index.html')


# ─── 扫描（预览检测结果）──────────────────────────────────────

@app.route('/api/scan', methods=['POST'])
def scan():
    """上传文件并扫描敏感信息，返回检测结果预览"""
    if 'file' not in request.files:
        return jsonify({'error': '未上传文件'}), 400

    f = request.files['file']
    ext = os.path.splitext(f.filename)[1].lower()
    if ext not in config.ALLOWED_EXTENSIONS:
        return jsonify({'error': f'不支持的文件格式: {ext}'}), 400

    # 保存上传文件（保留中文文件名）
    original_name = f.filename
    safe_name = f'{_ts()}_{_safe_filename(original_name)}'
    upload_path = os.path.join(config.UPLOAD_DIR, safe_name)
    f.save(upload_path)

    # 提取文本内容用于检测
    try:
        if ext in TEXT_EXTS:
            with open(upload_path, 'r', encoding='utf-8') as fh:
                content = fh.read()
        elif ext in DOCX_EXTS:
            from docx import Document
            doc = Document(upload_path)
            parts = []
            for para in doc.paragraphs:
                parts.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        parts.append(cell.text)
            content = '\n'.join(parts)
        elif ext in XLSX_EXTS:
            from openpyxl import load_workbook
            wb = load_workbook(upload_path)
            parts = []
            for ws in wb.worksheets:
                for row in ws.iter_rows():
                    for cell in row:
                        if cell.value and isinstance(cell.value, str):
                            parts.append(cell.value)
            content = '\n'.join(parts)
        else:
            return jsonify({'error': '不支持的文件格式'}), 400
    except Exception as e:
        return jsonify({'error': f'读取文件失败: {str(e)}'}), 500

    summary = detector.detect_summary(content)
    summary['filename'] = original_name
    summary['upload_name'] = safe_name
    summary['original_name'] = original_name  # 原始文件名，用于后续命名

    return jsonify(summary)


# ─── 脱敏 ──────────────────────────────────────────────────────

@app.route('/api/mask', methods=['POST'])
def mask():
    """执行脱敏，返回脱敏后的文件下载链接"""
    data = request.json
    upload_name = data.get('upload_name')
    if not upload_name:
        return jsonify({'error': '缺少文件名'}), 400

    upload_path = os.path.join(config.UPLOAD_DIR, upload_name)
    if not os.path.exists(upload_path):
        return jsonify({'error': '文件不存在'}), 404

    # 前端传来的已筛选检测项（用户可能删掉了误报项）
    items = data.get('items', None)
    # 前端传来的原始文件名（保留中文）
    original_name = data.get('original_name', '')

    ext = os.path.splitext(upload_name)[1].lower()
    # 用原始文件名命名输出，如果没传则从 upload_name 推断
    if original_name:
        base_name = os.path.splitext(original_name)[0]
    else:
        name_part = upload_name.split('_', 2)[-1] if '_' in upload_name else upload_name
        base_name = os.path.splitext(name_part)[0]
    ts = _ts()

    try:
        if ext in TEXT_EXTS:
            masked_content, mapping = mask_text_file(upload_path, detector, items)
            out_name = _safe_filename(f'{base_name}_脱敏{ext}')
            out_path = os.path.join(config.OUTPUT_DIR, out_name)
            save_text(masked_content, out_path)
        elif ext in DOCX_EXTS:
            doc, mapping = mask_docx_file(upload_path, detector, items)
            out_name = _safe_filename(f'{base_name}_脱敏{ext}')
            out_path = os.path.join(config.OUTPUT_DIR, out_name)
            doc.save(out_path)
        elif ext in XLSX_EXTS:
            wb, mapping = mask_xlsx_file(upload_path, detector, items)
            out_name = _safe_filename(f'{base_name}_脱敏{ext}')
            out_path = os.path.join(config.OUTPUT_DIR, out_name)
            wb.save(out_path)
        else:
            return jsonify({'error': '不支持的格式'}), 400
    except Exception as e:
        return jsonify({'error': f'脱敏失败: {str(e)}'}), 500

    # 保存加密映射表（命名：原始文件名_脱敏数量_时间戳.enc）
    safe_base = _safe_filename(base_name)
    mapping_name = f'{safe_base}_{len(mapping)}项_{ts}.enc'
    mapping_path = os.path.join(config.MAPPING_DIR, mapping_name)
    crypto.save_mapping(mapping, mapping_path)

    return jsonify({
        'output_file': out_name,
        'mapping_file': mapping_name,
        'masked_count': len(mapping),
    })


# ─── 还原 ──────────────────────────────────────────────────────

@app.route('/api/unmask', methods=['POST'])
def unmask():
    """执行还原"""
    if 'file' not in request.files:
        return jsonify({'error': '未上传文件'}), 400

    mapping_name = request.form.get('mapping')
    if not mapping_name:
        return jsonify({'error': '未选择映射表'}), 400

    mapping_path = os.path.join(config.MAPPING_DIR, mapping_name)
    if not os.path.exists(mapping_path):
        return jsonify({'error': '映射表不存在'}), 404

    f = request.files['file']
    ext = os.path.splitext(f.filename)[1].lower()
    if ext not in config.ALLOWED_EXTENSIONS:
        return jsonify({'error': f'不支持的文件格式: {ext}'}), 400

    # 保存上传文件（保留中文文件名）
    safe_name = f'{_ts()}_{_safe_filename(f.filename)}'
    upload_path = os.path.join(config.UPLOAD_DIR, safe_name)
    f.save(upload_path)

    # 解密映射表
    try:
        mapping = crypto.load_mapping(mapping_path)
    except Exception as e:
        return jsonify({'error': f'映射表解密失败: {str(e)}'}), 500

    original_name = f.filename
    base_name = os.path.splitext(original_name)[0]
    # 移除脱敏后缀
    for suffix in ('_脱敏', '_masked'):
        if base_name.endswith(suffix):
            base_name = base_name[:-len(suffix)]
            break

    try:
        if ext in TEXT_EXTS:
            restored = unmask_text_file(upload_path, mapping)
            out_name = _safe_filename(f'{base_name}_已还原{ext}')
            out_path = os.path.join(config.OUTPUT_DIR, out_name)
            save_text(restored, out_path)
        elif ext in DOCX_EXTS:
            doc = unmask_docx_file(upload_path, mapping)
            out_name = _safe_filename(f'{base_name}_已还原{ext}')
            out_path = os.path.join(config.OUTPUT_DIR, out_name)
            doc.save(out_path)
        elif ext in XLSX_EXTS:
            wb = unmask_xlsx_file(upload_path, mapping)
            out_name = _safe_filename(f'{base_name}_已还原{ext}')
            out_path = os.path.join(config.OUTPUT_DIR, out_name)
            wb.save(out_path)
        else:
            return jsonify({'error': '不支持的格式'}), 400
    except Exception as e:
        return jsonify({'error': f'还原失败: {str(e)}'}), 500

    return jsonify({
        'output_file': out_name,
    })


# ─── 文件下载 ──────────────────────────────────────────────────

@app.route('/api/download/<path:filename>')
def download(filename):
    """下载输出文件"""
    # 防止路径穿越
    filename = os.path.basename(filename)
    filepath = os.path.join(config.OUTPUT_DIR, filename)
    if not os.path.exists(filepath):
        return jsonify({'error': '文件不存在'}), 404
    return send_file(filepath, as_attachment=True)


# ─── 映射表管理 ────────────────────────────────────────────────

@app.route('/api/mappings')
def list_mappings():
    """列出所有映射表"""
    mappings = crypto.list_mappings(config.MAPPING_DIR)
    return jsonify(mappings)


@app.route('/api/mappings/<path:filename>', methods=['DELETE'])
def delete_mapping(filename):
    """删除映射表"""
    filename = os.path.basename(filename)
    filepath = os.path.join(config.MAPPING_DIR, filename)
    if os.path.exists(filepath):
        os.remove(filepath)
        return jsonify({'ok': True})
    return jsonify({'error': '文件不存在'}), 404


# ─── 自定义敏感词 ──────────────────────────────────────────────

@app.route('/api/custom-words', methods=['GET'])
def get_custom_words():
    """获取自定义敏感词列表"""
    words = []
    if os.path.exists(config.CUSTOM_NAMES_FILE):
        with open(config.CUSTOM_NAMES_FILE, 'r', encoding='utf-8') as f:
            for line in f:
                word = line.strip()
                if word and not word.startswith('#'):
                    words.append(word)
    return jsonify(words)


@app.route('/api/custom-words', methods=['POST'])
def save_custom_words():
    """保存自定义敏感词列表"""
    global detector
    words = request.json.get('words', [])
    with open(config.CUSTOM_NAMES_FILE, 'w', encoding='utf-8') as f:
        f.write('# 自定义敏感词列表（每行一个）\n')
        for word in words:
            if word.strip():
                f.write(word.strip() + '\n')
    # 重新加载检测器
    detector = SensitiveDetector(config.CUSTOM_NAMES_FILE)
    return jsonify({'ok': True, 'count': len(words)})


# ─── 启动 ─────────────────────────────────────────────────────

if __name__ == '__main__':
    import sys, io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8')
    print(f'\n[*] 敏感信息加密工具已启动')
    print(f'[*] 打开浏览器访问: http://localhost:{config.PORT}')
    print(f'[*] 输出目录: {config.OUTPUT_DIR}')
    print(f'[*] 映射表目录: {config.MAPPING_DIR}\n')
    app.run(host='0.0.0.0', port=config.PORT, debug=True)
